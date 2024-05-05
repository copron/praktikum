import os
import shutil
from docx import Document

def create_directory(directory_name):
    try:
        os.mkdir(directory_name)
        print(f"Директория '{directory_name}' успешно создана.")
    except FileExistsError:
        print(f"Директория '{directory_name}' уже существует.")

def delete_directory(directory_name):
    try:
        os.rmdir(directory_name)
        print(f"Директория '{directory_name}' успешно удалена.")
    except FileNotFoundError:
        print(f"Директория '{directory_name}' не найдена.")
    except OSError as e:
        if '[WinError 145]' in str(e):  # если в ошибке [WinError 145]
            choice = input(f"Директория '{directory_name}' не пуста. Вы уверены, что хотите удалить её? (да/нет): ")
            if choice.lower() == 'да':
                shutil.rmtree(directory_name)
                print(f"Директория '{directory_name}' успешно удалена вместе со всем её содержимым.")
            else:
                print("Удаление отменено.")
        else:
            print(f"Ошибка при удалении директории '{directory_name}': {e}")


def list_contents():
    contents = os.listdir()
    print("Содержимое текущей директории:")
    for item in contents:
        print(item)

def change_directory(new_directory):
    try:
        os.chdir(new_directory)
        print(f"Текущая директория изменена на '{new_directory}'.")
    except FileNotFoundError:
        print(f"Директория '{new_directory}' не найдена.")

def create_file(file_name, content=''):
    try:
        if file_name.endswith('.docx'):
            doc = Document()
            for line in content.split('\n'):
                doc.add_paragraph(line)
            doc.save(file_name)
            print(f"Файл '{file_name}' успешно создан и заполнен.")
        else:
            with open(file_name, 'w') as file:
                file.write(content)
            print(f"Файл '{file_name}' успешно создан.")
    except FileExistsError:
        print(f"Файл '{file_name}' уже существует.")
    except Exception as e:
        print(f"Ошибка при создании файла '{file_name}': {e}")


def read_file(file_name):
    try:
        if file_name.endswith('.docx'):
            doc = Document(file_name)
            content = ""
            for para in doc.paragraphs:
                content += para.text + "\n"
        else:
            with open(file_name, 'r') as file:
                content = file.read()
        print(f"Содержимое файла '{file_name}':\n{content}")
    except FileNotFoundError:
        print(f"Файл '{file_name}' не найден.")
    except Exception as e:
        print(f"Ошибка при чтении файла '{file_name}': {e}")

def write_file(file_name, content):
    try:
        if file_name.endswith('.docx'):
            doc = Document()
            for line in content.split('\n'):
                doc.add_paragraph(line)
            doc.save(file_name)
        else:
            with open(file_name, 'w') as file:
                file.write(content)
        print(f"Содержимое файла '{file_name}' успешно обновлено.")
    except FileNotFoundError:
        print(f"Файл '{file_name}' не найден.")
    except Exception as e:
        print(f"Ошибка при записи в файл '{file_name}': {e}")

def delete_file(file_name):
    try:
        os.remove(file_name)
        print(f"Файл '{file_name}' успешно удален.")
    except FileNotFoundError:
        print(f"Файл '{file_name}' не найден.")
    except Exception as e:
        print(f"Ошибка при удалении файла '{file_name}': {e}")

def copy_file(source_file, destination):
    try:
        shutil.copy(source_file, destination)
        print(f"Файл '{source_file}' успешно скопирован в '{destination}'.")
    except FileNotFoundError:
        print(f"Файл '{source_file}' не найден.")
    except Exception as e:
        print(f"Ошибка при копировании файла '{source_file}': {e}")

def move_file(source_file, destination):
    try:
        shutil.move(source_file, destination)
        print(f"Файл '{source_file}' успешно перемещен в '{destination}'.")
    except FileNotFoundError:
        print(f"Файл '{source_file}' не найден.")
    except Exception as e:
        print(f"Ошибка при перемещении файла '{source_file}': {e}")

def rename_file(old_name, new_name):
    try:
        os.rename(old_name, new_name)
        print(f"Файл '{old_name}' успешно переименован в '{new_name}'.")
    except FileNotFoundError:
        print(f"Файл '{old_name}' не найден.")
    except Exception as e:
        print(f"Ошибка при переименовании файла '{old_name}': {e}")
